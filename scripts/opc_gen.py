from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
st = doc.styles["Normal"]
st.font.size = Pt(11)
st.paragraph_format.line_spacing = 1.5

def gh(text, lv=1):
    h = doc.add_heading(text, level=lv)
    for r in h.runs: r.font.color.rgb = RGBColor(7,193,96)

def at(hd, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(hd), style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hd): t.rows[0].cells[i].text = h
    for ri,row in enumerate(rows,1):
        for ci,v in enumerate(row): t.rows[ri].cells[ci].text = str(v)

for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MediSlim OPC 综合运营手册'); r.font.size = Pt(48); r.font.bold = True; r.font.color.rgb = RGBColor(7,193,96)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('一人公司+AI驱动的消费医疗平台
完整操作指南 v1.0'); r.font.size = Pt(20); r.font.color.rgb = RGBColor(100,100,100)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('日期：2026-04-04 编制：贾维斯AI系统'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(150,150,150)
doc.add_page_break()

gh('目  录')
for t in ['一、项目概述与OPC定位','二、AI四层架构','三、人员配置与阶段规划','四、AI自动化清单','五、三圈产品矩阵详解','六、业务流程详解','七、全渠道运营SOP','八、资金流与财务流程','九、数据看板与决策体系','十、合规检查清单','十一、30天执行计划','十二、应急手册']:
    doc.add_paragraph(t)
doc.add_page_break()

gh('一、项目概述与OPC定位')
doc.add_paragraph('MediSlim是一个AI驱动的消费医疗+健康消费品平台，借鉴美国Medvi公司(年营收4亿美金)的极简架构模式，结合中国特色打造从专业医疗到日常养生的全覆盖健康消费平台。')
doc.add_paragraph('OPC(One Person Company)模式核心：AI驱动90%日常运营，创始人只做关键决策。MIMO API零成本无限推理，三圈产品矩阵覆盖全消费场景，全外包履约消除重资产。')
doc.add_paragraph('我们不是医疗机构，是流量+技术服务公司。处方权归合作医师，药品销售归合作药房，我们做品牌、流量、AI评估、患者管理、复购运营。')
at(['维度','Medvi(美国)','MediSlim(中国)'],[['品类','单一(GLP-1)','全品类(处方+药食+保健品)'],['文化','无文化认同','中医体质辨识+药食同源'],['合规','FDA灰色地带','完全合规(持牌合作)'],['护城河','零','AI体质辨识+用户数据'],['利润率','16.2%','预估35%'],['AI能力','付费API','MIMO API(无限额度零成本)']])
doc.add_page_break()

gh('二、AI四层架构(OPC核心引擎)')
doc.add_paragraph('第一层：AI Agent层(24/7自动运转)')
for i in ['内容生成Agent：自动生成小红书/抖音/公众号内容','客服Agent：自动回复+问答+分流','数据分析Agent：自动报表+预警+ROI','运营Agent：自动复购提醒+社群任务']:
    doc.add_paragraph(i, style='List Bullet')
doc.add_paragraph('第二层：工具层(已建设)')
for i in ['app.py :8090 产品评估+下单系统','landing.py :8091 获客落地页','flow_engine.py :8092 15状态业务流引擎','admin.py :8093 管理后台']:
    doc.add_paragraph(i, style='List Bullet')
doc.add_paragraph('第三层：合作层(全外包)')
at(['类型','候选方','作用'],[['医疗层','微医/好大夫/京东健康','处方开方'],['药房层','大参林/益丰/老百姓','药品销售+配送'],['物流层','顺丰/京东','全品类配送'],['生产层','SC食品工厂/保健品代工','药食同源/保健品生产'],['支付层','微信支付/支付宝','收款+退款']])
doc.add_paragraph('第四层：人层(小林医生)：决策/合规/关键节点/AI监督')
doc.add_page_break()

gh('三、人员配置与阶段规划')
at(['阶段','时间','人数','分工'],[['阶段1','M1-M3','1人','小林医生=CEO+产品+合规；AI全包其余'],['阶段2','M4-M6','2人','+1运营助理(客服/社群/发货)'],['阶段3','M7-M12','3-5人','+1内容(视频/直播)+1技术(系统维护)']])
doc.add_paragraph()
doc.add_paragraph('OPC精简逻辑：AI替代整个部门。内容团队→AI Agent；客服→AI Agent；数据分析→AI Agent。人只做AI做不了的事：决策、谈判、出镜、投诉处理。')
doc.add_page_break()

gh('四、AI自动化清单')
at(['任务','执行者','频率','AI能力'],[['写小红书笔记','AI','日更3-5条','全自动'],['写抖音脚本','AI','日更1-2条','全自动'],['写公众号文章','AI','周更2篇','全自动'],['用户体质评估','AI','实时','全自动'],['产品推荐','AI','实时','全自动'],['禁忌筛查','AI','实时','全自动'],['复购提醒','AI','按周期','全自动'],['数据报表','AI','日报/周报','全自动'],['客户首次咨询','AI','实时','全自动'],['复杂咨询/投诉','人','不定期','需人工'],['合作方谈判','人','按需','需人工'],['资质申请','人','按需','需人工'],['视频出镜','人+AI','按需','AI辅助']])
doc.add_page_break()

gh('五、三圈产品矩阵详解')
doc.add_heading('5.1 第一圈：处方医疗(引流+信任)', level=2)
doc.add_paragraph('需要互联网医院医师开方，建立专业品牌形象。处方药是引流入口。')
at(['品类','产品','首月价','续费/月','市场规模','合规'],[['GLP-1减重','司美格鲁肽/替尔泊肽','399','599','500亿+','处方药+互联网医院'],['防脱生发','米诺地尔+非那雄胺','199','299','200亿+','处方药'],['男性健康','西地那非/他达拉非','299','499','150亿+','处方药'],['助眠','褪黑素+处方药','199','299','100亿+','部分OTC'],['皮肤处方','维A酸/阿达帕林','249','349','80亿+','处方药']])
doc.add_heading('5.2 第二圈：药食同源(复购+文化认同)', level=2)
doc.add_paragraph('核心差异化竞争力。食品级产品，合规门槛极低(SC食品许可证即可)，文化认同度极高，复购率80%+，毛利率60-70%。')
at(['品类','产品配方','月价','复购率','市场规模'],[['祛湿轻体','红豆薏米+茯苓+陈皮','168','80%','300亿+'],['气血调养','红枣+枸杞+阿胶+黄芪','198','85%','400亿+'],['护肝养生','葛根+枳椇子+菊花','168','75%','200亿+'],['暖宫驱寒','姜+红糖+桂圆+当归','158','80%','150亿+'],['健脾养胃','山药+莲子+芡实+薏米','148','78%','180亿+'],['润肺清燥','雪梨+百合+银耳+枇杷','138','70%','100亿+']])
doc.add_heading('5.3 第三圈：保健品(利润+规模)', level=2)
at(['品类','产品','月价','复购率','目标人群'],[['益生菌','多菌株复合益生菌','198','80%','肠胃/减重辅助'],['胶原蛋白','鱼胶原蛋白肽','228','75%','女性美容'],['维生素D','VD3+K2复合','98','70%','办公室人群'],['Omega-3','高纯度鱼油','168','70%','心血管/大脑'],['辅酶Q10','泛醇CoQ10','188','68%','心脏/抗衰'],['叶黄素','叶黄素酯+玉米黄质','128','72%','护眼'],['蛋白粉','乳清蛋白/植物蛋白','188','75%','运动/减脂'],['综合维生素','男士/女士多维','128','80%','全人群']])
doc.add_heading('5.4 产品接入标准化SOP', level=2)
for s,d in [('Step 1 合规审核','是否需处方/广告批文/特殊资质'),('Step 2 供应链搭建','找供应商+签合同+定价'),('Step 3 系统接入','产品录入/评估问卷/定价/物流配置'),('Step 4 内容准备','AI生成50条小红书+20条抖音+5篇公众号'),('Step 5 上线运营','内测→正式上线→建社群→复购循环')]:
    doc.add_paragraph(s+'：'+d)
doc.add_heading('5.5 各品类接入时间表', level=2)
at(['启动周','品类','依赖','周期'],[['W1-2','祛湿轻体/气血调养','SC工厂','2周'],['W2-3','益生菌/胶原蛋白','供应商对接','2周'],['W3-4','GLP-1减重/防脱','互联网医院+药房','4周'],['W4-5','男性健康/助眠','互联网医院+药房','3周'],['W5-6','皮肤处方','互联网医院+药房','3周']])
doc.add_page_break()

gh('六、业务流程详解')
doc.add_heading('6.1 完整用户旅程(9步)', level=2)
for i,(s,d) in enumerate([('触达','社交媒体广告/社群推荐/口碑传播'),('落地','H5落地页/企微社群/朋友推荐'),('评估','AI健康问卷5分钟：处方类→症状+禁忌 | 药食同源→体质辨识 | 保健品→营养需求'),('方案','个性化健康方案(AI生成+医师审核)'),('下单','微信支付/支付宝付款'),('履约','处方→互联网医院开方→药房→顺丰 | 药食同源/保健品→工厂/仓→顺丰'),('随访','企微1对1跟踪+每周打卡+AI调方案'),('复购','Day 21提醒→一键续费→自动发货→循环'),('裂变','邀请好友→双方优惠→社群运营→滚雪球')],1):
    doc.add_paragraph(f'Step {i} {s}：{d}')
doc.add_heading('6.2 订单状态机(15个状态)', level=2)
doc.add_paragraph('new_lead → contacted → assessed → paid → ih_submitted → doctor_review → prescribed → pharmacy_order → dispensing → shipped → delivered → in_use → refill_reminder → refill_paid → cancelled')
doc.add_paragraph('每个状态都有超时机制和自动推进逻辑。从线索到用药最快2天，全自动流转。')
doc.add_heading('6.3 AI体质辨识', level=2)
doc.add_paragraph('15道中医体质辨识题，AI判断九种体质，推荐对应药食同源方案。')
for t in ['平和质→通用保健','气虚质→黄芪+党参+山药','阳虚质→肉桂+干姜+韭菜子','阴虚质→百合+枸杞+石斛','痰湿质→薏米+茯苓+陈皮(最常见)','湿热质→菊花+薏米+绿豆','血瘀质→山楂+红花+丹参','气郁质→玫瑰+佛手+香橼','特禀质→过敏体质调理']:
    doc.add_paragraph(t, style='List Bullet')
doc.add_page_break()

gh('七、全渠道运营SOP')
doc.add_heading('7.1 企业微信(立即启动)', level=2)
doc.add_paragraph('客服SOP(Day 0到Day 28)：')
at(['时间','动作','内容'],[['Day 0','加好友','欢迎语+AI评估链接'],['Day 1','推报告','AI评估报告+推荐产品'],['Day 3','跟反馈','使用感受+解答疑问'],['Day 7','邀入群','产品群+首购优惠'],['Day 14','效果跟进','打卡反馈+效果记录'],['Day 21','复购提醒','一键续费链接'],['Day 28','续费优惠','续费折扣+升级建议']])
doc.add_paragraph('社群运营：每日早报(养生知识)+每周直播(医师答疑)+每月活动(限时优惠)+用户故事分享')
doc.add_heading('7.2 微信公众号(4月8日后启动)', level=2)
doc.add_paragraph('自动回复：关注→欢迎语+AI评估链接 | 关键词→产品介绍 | 其他→引导加企微')
doc.add_paragraph('底部菜单：健康评估→H5 | 我的订单→查询 | 咨询顾问→企微')
doc.add_heading('7.3 小红书(立即启动)', level=2)
doc.add_paragraph('账号矩阵：1个官方号+3个人设号(减重/养生/医生科普)+10个素人号')
doc.add_paragraph('内容节奏：早8点养生知识 / 中12点案例产品 / 晚8点科普互动')
doc.add_paragraph('爆文公式：标题=数字+痛点+方案 | 封面=对比图或数据图 | 正文=故事+干货+CTA')
doc.add_heading('7.4 抖音(立即启动)', level=2)
doc.add_paragraph('账号：1官方号+1达人号')
doc.add_paragraph('短视频公式：前3秒钩子→中间痛点+方案+数据→结尾CTA(评论区扣1/关注)')
doc.add_page_break()

gh('八、资金流与财务流程')
doc.add_paragraph('收入流向：用户付款→微信/支付宝商户号→分流到各环节')
at(['流向','占比','说明'],[['平台留存(利润)','35%','我们的净利润'],['互联网医院','10%','处方费分成'],['药房','35%','药品成本'],['SC工厂/供应商','20%','药食同源/保健品成本'],['顺丰/物流','5%','配送费'],['获客/广告','15%','流量投放']])
doc.add_paragraph('现金流特点：预收后付(用户先付款，T+30给供应商)，首月即可正现金流。订阅制自动续费，现金流可预测。')
doc.add_page_break()

gh('九、数据看板与决策体系')
doc.add_heading('9.1 核心5指标(每日必看)', level=2)
for i in ['今日新增线索数','今日新订单数','今日收入(¥)','线索→订单转化率','本月MRR(月经常性收入)']:
    doc.add_paragraph(i, style='List Bullet')
doc.add_heading('9.2 渠道指标', level=2)
doc.add_paragraph('各渠道CAC(获客成本) / LTV(生命周期价值) / ROI / 各品类复购率')
doc.add_heading('9.3 预警指标', level=2)
at(['预警','阈值','动作'],[['未复购','7天未复购','AI自动提醒'],['投诉率','>5%','人工介入排查'],['退货率','>10%','暂停该品类'],['获客成本','CAC>LTV','停止该渠道投放']])
doc.add_page_break()

gh('十、合规检查清单')
doc.add_heading('10.1 公司层面', level=2)
for i in ['营业执照(经营范围含健康咨询+食品销售+信息技术)','食品经营许可证','增值电信业务经营许可证(ICP)']:
    doc.add_paragraph(i, style='List Bullet')
doc.add_heading('10.2 产品层面', level=2)
for i in ['处方药：必须通过互联网医院，我们只做信息撮合','药食同源：SC食品许可证+标签合规','保健品：蓝帽子批文(如有)+广告审查']:
    doc.add_paragraph(i, style='List Bullet')
doc.add_heading('10.3 宣传层面', level=2)
for i in ['不使用禁词(根治/最佳/唯一/国家级)','不夸大功效暗示疗效','使用真实案例+用户授权','广告内容法务审核']:
    doc.add_paragraph(i, style='List Bullet')
doc.add_page_break()

gh('十一、30天执行计划')
at(['时间','任务','产出'],[['Day 1-3','文档+方案','商业方案/业务流/管理后台/OPC手册'],['Day 4-7','供应链调研','3家SC工厂/3家保健品/2家互联网医院/顺丰签约'],['Day 8-10','系统开发','体质辨识AI/产品详情页/支付接入'],['Day 11-14','内容准备','50条小红书/20条抖音/5篇公众号/企微配置'],['Day 15-17','内测启动','邀请100个种子用户'],['Day 18-21','反馈优化','收集问题/优化流程/修复bug'],['Day 22-25','正式上线','公众号菜单/开始投放/建社群'],['Day 26-30','运营启动','首批付费用户/复购循环/数据追踪']])
doc.add_page_break()

gh('十二、应急手册')
at(['风险','概率','影响','应急方案'],[['互联网医院终止合作','中','高','立即切换备选/暂停处方订单/通知用户'],['供应中断','高','中','启用备选供应商/调整产品组合/通知+补偿'],['微信封号','中','高','多平台迁移/备份用户数据/企微继续'],['大规模投诉','低','高','暂停新订单/逐排查/退费补偿/公开声明'],['政策收紧','低','高','合规底线优先/随时调整品类']])
doc.add_page_break()

gh('附录')
doc.add_paragraph('项目仓库：https://github.com/MoKangMedical/medi-slim')
doc.add_paragraph('管理后台：http://localhost:8093')
doc.add_paragraph('已上线服务：app.py(:8090) / landing.py(:8091) / flow_engine.py(:8092) / admin.py(:8093)')
doc.add_paragraph()
at(['版本','日期','编制','内容'],[['v1.0','2026-04-04','贾维斯(Jarvis)','初始版本，12章完整OPC运营手册']])

path = '/root/medi-slim/docs/MediSlim-OPC运营手册-v1.0.docx'
doc.save(path)
print(f'DONE {os.path.getsize(path)//1024}KB')
