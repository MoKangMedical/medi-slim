"""
生成 MediSlim 正式商业方案 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.color.rgb = RGBColor(7, 193, 96)  # 绿色

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MediSlim')
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = RGBColor(7, 193, 96)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('中国版Medvi消费医疗平台\n完整商业方案 v1.0')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('日期：2026年4月4日\n编制：贾维斯（Jarvis）AI系统\n版本：v1.0\n状态：内部文件')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_page_break()

# ========== 目录 ==========
doc.add_heading('目  录', level=1)
toc_items = [
    '一、项目概述',
    '二、产品矩阵（三圈模型）',
    '三、业务流程',
    '四、合规架构与风控',
    '五、财务模型',
    '六、获客与营销策略',
    '七、技术架构',
    '八、30天启动计划',
    '九、竞争优势分析',
    '十、风险提示与应急预案',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ========== 正文 ==========

# --- 一、项目概述 ---
doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 项目定位', level=2)
doc.add_paragraph(
    'MediSlim是一个AI驱动的消费医疗+健康消费品平台，借鉴美国Medvi公司（年营收$4亿，16.2%净利润率）的'
    '"极简架构"模式，结合中国特色（药食同源文化、互联网医疗合规框架、微信生态），'
    '打造从专业医疗到日常养生的全覆盖健康消费平台。'
)

doc.add_heading('1.2 核心逻辑', level=2)
doc.add_paragraph(
    '处方药（引流+信任建立）→ 药食同源（高频复购+文化认同）→ 保健品（高毛利+规模化）\n\n'
    '三圈飞轮效应：处方药带来精准用户 → AI体质辨识建立信任 → 药食同源产品实现高频复购 → '
    '保健品提升客单价和利润 → 用户数据积累优化AI推荐 → 滚雪球增长。'
)

doc.add_heading('1.3 我们的定位', level=2)
p = doc.add_paragraph()
p.add_run('我们不是医疗机构，我们是"流量+技术服务"公司：\n').bold = True
items = [
    '处方权 → 合作互联网医院持有',
    '药品销售 → 合作药房持有',
    '我们做 → 品牌、流量、AI评估、患者管理、复购运营',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# --- 二、产品矩阵 ---
doc.add_heading('二、产品矩阵（三圈模型）', level=1)

doc.add_heading('2.1 第一圈：处方医疗（引流+信任）', level=2)
doc.add_paragraph('需要互联网医院医师开方，建立"专业"品牌形象。')

table = doc.add_table(rows=6, cols=6, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['品类', '产品', '首月价', '续费/月', '市场规模', '合规要求']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['GLP-1减重', '司美格鲁肽/替尔泊肽', '¥399', '¥599', '500亿+', '处方药+互联网医院'],
    ['防脱生发', '米诺地尔+非那雄胺', '¥199', '¥299', '200亿+', '处方药(非那雄胺)'],
    ['男性健康', '西地那非/他达拉非', '¥299', '¥499', '150亿+', '处方药'],
    ['助眠', '褪黑素+必要时处方药', '¥199', '¥299', '100亿+', '部分OTC'],
    ['皮肤处方', '维A酸/阿达帕林', '¥249', '¥349', '80亿+', '处方药'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph()

doc.add_heading('2.2 第二圈：药食同源（复购+文化认同）', level=2)
doc.add_paragraph(
    '符合中国传统养生文化，无需处方，复购率极高。'
    '这是Medvi模式中完全没有的中国特色品类，也是核心差异化竞争力。'
)

table = doc.add_table(rows=7, cols=6, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['品类', '产品', '月价', '复购率', '市场规模', '优势']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['祛湿轻体', '红豆薏米+茯苓+陈皮', '¥168', '80%', '300亿+', '中国人90%认为自己湿气重'],
    ['气血调养', '红枣+枸杞+阿胶+黄芪', '¥198', '85%', '400亿+', '女性刚需，送礼佳品'],
    ['护肝养生', '葛根+枳椇子+菊花', '¥168', '75%', '200亿+', '应酬人群刚需'],
    ['暖宫驱寒', '姜+红糖+桂圆+当归', '¥158', '80%', '150亿+', '宫寒痛经女性'],
    ['健脾养胃', '山药+莲子+芡实+薏米', '¥148', '78%', '180亿+', '脾胃虚弱人群'],
    ['润肺清燥', '雪梨+百合+银耳+枇杷', '¥138', '70%', '100亿+', '秋冬季节性爆品'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('药食同源产品特点：')
run.bold = True
items = [
    '食品级产品，合规门槛极低（SC食品生产许可证即可）',
    '文化认同度极高——"老祖宗的智慧"',
    '复购率80%+（每月定期配送订阅制）',
    '毛利率60-70%（原料成本极低）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 第三圈：保健品/营养补充（利润+规模）', level=2)

table = doc.add_table(rows=9, cols=5, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['品类', '产品', '月价', '复购率', '目标人群']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['益生菌', '多菌株复合益生菌', '¥198', '80%', '肠胃调理/减重辅助'],
    ['胶原蛋白', '鱼胶原蛋白肽', '¥228', '75%', '女性美容'],
    ['维生素D', 'VD3+K2复合', '¥98', '70%', '办公室人群'],
    ['Omega-3', '高纯度鱼油', '¥168', '70%', '心血管/大脑健康'],
    ['辅酶Q10', '泛醇CoQ10', '¥188', '68%', '心脏/抗衰'],
    ['叶黄素', '叶黄素酯+玉米黄质', '¥128', '72%', '护眼(手机族)'],
    ['蛋白粉', '乳清蛋白/植物蛋白', '¥188', '75%', '运动/减脂'],
    ['综合维生素', '男士/女士多维', '¥128', '80%', '全人群'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_page_break()

# --- 三、业务流程 ---
doc.add_heading('三、业务流程', level=1)

doc.add_heading('3.1 完整用户旅程', level=2)
steps = [
    ('触达', '社交媒体广告 / 社群推荐 / 口碑传播'),
    ('落地', 'H5落地页 → 企微社群 → 朋友推荐'),
    ('评估', 'AI健康问卷（5分钟）：处方类→症状评估+禁忌筛查 | 药食同源→体质辨识 | 保健品→营养需求评估'),
    ('方案', '个性化健康方案（AI生成+医师审核）'),
    ('下单', '微信支付/支付宝付款'),
    ('履约', '处方类→互联网医院开方→药房配药→顺丰到家 | 药食同源/保健品→工厂/仓库直发'),
    ('随访', '企微健康顾问1对1跟踪 + 每周打卡反馈 + AI自动调方案'),
    ('复购', 'Day 21自动提醒 → 一键续费 → 自动发货 → 循环'),
    ('裂变', '邀请好友 → 双方优惠 → 社群运营 → 滚雪球'),
]
for i, (step, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}：{step}')
    run.bold = True
    doc.add_paragraph(desc)

doc.add_heading('3.2 订单状态机', level=2)
doc.add_paragraph(
    '系统包含15个订单状态，全自动流转：\n\n'
    '新线索 → 已联系 → 已评估 → 已付费 → 互联网医院提交 → 医师审核 → 已开处方 → '
    '药房下单 → 配药中 → 已发货 → 已签收 → 用药中 → 续费提醒 → 已续费(循环) → 已取消(终态)\n\n'
    '每个状态都有超时机制和自动推进逻辑，确保订单不被卡住。'
)

doc.add_heading('3.3 AI体质辨识（中国特色）', level=2)
doc.add_paragraph(
    '这是Medvi完全不具备的中国特色能力。通过15道中医体质辨识题，'
    'AI判断用户属于九种体质中的哪一种，推荐对应的药食同源方案：'
)
types = [
    '平和质 → 通用保健方案',
    '气虚质 → 补气方案（黄芪+党参+山药）',
    '阳虚质 → 温阳方案（肉桂+干姜+韭菜子）',
    '阴虚质 → 滋阴方案（百合+枸杞+石斛）',
    '痰湿质 → 祛湿方案（薏米+茯苓+陈皮）★ 最常见',
    '湿热质 → 清热方案（菊花+薏米+绿豆）',
    '血瘀质 → 活血方案（山楂+红花+丹参）',
    '气郁质 → 疏肝方案（玫瑰+佛手+香橼）',
    '特禀质 → 过敏体质调理',
]
for t in types:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# --- 四、合规架构 ---
doc.add_heading('四、合规架构与风控', level=1)

doc.add_heading('4.1 法规依据', level=2)
table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['法规', '适用范围', '我们的合规点']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['《互联网诊疗管理办法》', '处方药销售', '依托持牌互联网医院'],
    ['《药品管理法》', '药品经营', '不做药品经营，只做信息撮合'],
    ['《食品安全法》', '药食同源产品', 'SC食品许可证+合规标签'],
    ['《广告法》', '产品宣传', '不使用禁词，不夸大功效'],
    ['《个人信息保护法》', '用户数据', '境内存储，加密处理'],
    ['《反不正当竞争法》', '市场推广', '真实案例，不虚构效果'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_heading('4.2 风控矩阵', level=2)
table = doc.add_table(rows=8, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['风险', '等级', '缓解措施', '责任方']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['无牌行医', '致命', '只做信息撮合，处方权归合作医师', '我们+互联网医院'],
    ['假药风险', '致命', '只对接获批药品，不碰合成药', '药房'],
    ['数据泄露', '致命', 'AES-256加密+境内存储+最小权限', '我们'],
    ['广告违规', '重大', '法务审核+禁词过滤+真实案例', '我们'],
    ['效果投诉', '重大', '退费机制+医师随访+数据存档', '我们+医师'],
    ['供应中断', '重大', '2-3家供应商并行', '我们'],
    ['微信封号', '重大', '合规内容+先H5后小程序', '我们'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_page_break()

# --- 五、财务模型 ---
doc.add_heading('五、财务模型', level=1)

doc.add_heading('5.1 单位经济模型', level=2)
p = doc.add_paragraph()
run = p.add_run('处方药单（GLP-1减重）：')
run.bold = True
doc.add_paragraph(
    '收入：¥399（首月）→ ¥599/月（续费）\n'
    '成本：药品¥180 + 互联网医院分成¥40 + 物流¥15 + 获客¥80 + 运营¥30 = ¥345\n'
    '首月毛利：¥54（13.5%）\n'
    '续费毛利：¥254/月（42.4%）\n'
    'LTV（12个月）：¥5,191，利润率约35%'
)

p = doc.add_paragraph()
run = p.add_run('药食同源单（祛湿轻体）：')
run.bold = True
doc.add_paragraph(
    '收入：¥168/月\n'
    '成本：原料¥35 + 包装¥10 + 物流¥12 + 获客¥40 = ¥97\n'
    '毛利：¥71/月（42.3%）\n'
    'LTV（12个月）：¥1,512，利润率约50%'
)

doc.add_heading('5.2 收入预测（12个月）', level=2)
table = doc.add_table(rows=7, cols=6, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['月份', '付费用户', '处方药收入', '药食同源收入', '保健品收入', '月总收入']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['M1', '100', '¥39,900', '¥16,800', '¥19,800', '¥76,500'],
    ['M2', '250', '¥99,750', '¥42,000', '¥49,500', '¥191,250'],
    ['M3', '500', '¥199,500', '¥84,000', '¥99,000', '¥382,500'],
    ['M6', '2,000', '¥798,000', '¥336,000', '¥396,000', '¥1,530,000'],
    ['M9', '5,000', '¥1,995,000', '¥840,000', '¥990,000', '¥3,825,000'],
    ['M12', '10,000', '¥3,990,000', '¥1,680,000', '¥1,980,000', '¥7,650,000'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Y1预估总收入：约¥3,000万 | Y1预估净利润率：25-30%')
run.bold = True
run.font.size = Pt(13)

doc.add_page_break()

# --- 六、获客策略 ---
doc.add_heading('六、获客与营销策略', level=1)

doc.add_heading('6.1 流量矩阵', level=2)
table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['渠道', '策略', '预算占比', '预期ROI']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['小红书', '素人种草+医生科普+案例分享', '30%', '1:5'],
    ['抖音', '短视频+直播+信息流', '25%', '1:4'],
    ['微信生态', '公众号+社群+朋友圈广告', '25%', '1:6'],
    ['私域', '企微1对1+社群+裂变', '15%', '1:8'],
    ['线下', '药房合作+社区推广', '5%', '1:3'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_heading('6.2 裂变机制', level=2)
items = [
    '邀请1人 → 双方各得¥30优惠券',
    '邀请3人 → 免费领取1个月药食同源产品',
    '邀请10人 → 成为"健康大使"，享8折+佣金10%',
    '邀请50人 → 区域合伙人，享7折+佣金15%',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# --- 七、技术架构 ---
doc.add_heading('七、技术架构', level=1)

doc.add_heading('7.1 系统组成', level=2)
doc.add_paragraph(
    '已上线服务：\n'
    '• app.py (端口8090) — 产品评估+下单系统\n'
    '• landing.py (端口8091) — 微信生态获客落地页\n'
    '• flow_engine.py (端口8092) — 完整业务流引擎（15个状态机）\n\n'
    '待开发：\n'
    '• 微信小程序前端\n'
    '• AI体质辨识模块\n'
    '• 订阅管理系统\n'
    '• 裂变系统\n'
    '• 互联网医院/药房/支付 API对接'
)

doc.add_heading('7.2 AI能力（MIMO API驱动）', level=2)
items = [
    '体质辨识：15题问卷 → AI九种体质判断',
    '产品推荐：基于体质+症状+偏好 → 个性化方案',
    '禁忌筛查：药物相互作用+过敏+孕期筛查',
    '随访提醒：AI自动判断复购时机+效果评估',
    '内容生成：批量生成小红书/抖音/公众号内容',
    '客服机器人：7×24智能问答',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# --- 八、30天启动计划 ---
doc.add_heading('八、30天启动计划', level=1)

doc.add_heading('Week 1：产品+系统', level=2)
items = [
    '✓ 商业方案（已完成）',
    '✓ 评估系统（已完成）',
    '✓ 落地页（已完成）',
    '✓ 业务流引擎（已完成）',
    '□ 药食同源产品线扩充',
    '□ 体质辨识AI开发',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Week 2：供应链+合规', level=2)
items = [
    '□ 联系2-3家互联网医院洽谈合作',
    '□ 药食同源SC食品供应商对接',
    '□ 保健品供应商筛选',
    '□ 公司注册/资质办理',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Week 3：内容+流量', level=2)
items = [
    '□ 小红书账号矩阵（10个号）',
    '□ 抖音账号注册+首批内容',
    '□ 微信公众号注册',
    '□ 企微社群搭建',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Week 4：内测+上线', level=2)
items = [
    '□ 邀请100人内测',
    '□ 修复问题+优化流程',
    '□ 正式上线',
    '□ 开始投放获客',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# --- 九、竞争优势 ---
doc.add_heading('九、竞争优势分析（vs Medvi）', level=1)

table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['维度', 'Medvi（美国）', 'MediSlim（中国）']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['品类', '单一（GLP-1）', '全品类（处方+药食+保健品）'],
    ['文化', '无文化认同', '中医体质辨识+药食同源'],
    ['合规', 'FDA灰色地带', '完全合规（持牌合作）'],
    ['护城河', '零', 'AI体质辨识+用户健康数据'],
    ['复购驱动', '药物依赖', '文化认同+习惯养成'],
    ['利润率', '16.2%', '预估35%'],
    ['AI能力', 'ChatGPT+Claude（付费）', 'MIMO API（无限额度零成本）'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_page_break()

# --- 十、风险提示 ---
doc.add_heading('十、风险提示与应急预案', level=1)

table = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['风险', '概率', '影响', '应急方案']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['互联网医院终止合作', '中', '高', '同时对接3家，不押注单一家'],
    ['GLP-1供应紧张', '高', '中', '药食同源产品不依赖进口原料'],
    ['微信封号', '中', '高', '多平台布局（抖音/小红书/APP）'],
    ['竞争对手复制', '高', '中', '快速迭代+数据壁垒+品牌建设'],
    ['政策收紧', '低', '高', '合规底线优先，随时调整品类'],
    ['用户大规模投诉', '低', '高', '退费机制+医师随访+数据存档'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

# ========== 附录 ==========
doc.add_page_break()
doc.add_heading('附录', level=1)

doc.add_heading('附录A：项目仓库', level=2)
doc.add_paragraph(
    'MediSlim仓库：https://github.com/MoKangMedical/medi-slim\n'
    'MediChat-RD仓库：https://github.com/MoKangMedical/medichat-rd\n\n'
    '本地服务：\n'
    '• 评估系统：http://localhost:8090\n'
    '• 落地页：http://localhost:8091\n'
    '• 业务流引擎：http://localhost:8092'
)

doc.add_heading('附录B：文档版本历史', level=2)
table = doc.add_table(rows=2, cols=4, style='Light Grid Accent 1')
headers = ['版本', '日期', '编制', '内容']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = ['v1.0', '2026-04-04', '贾维斯（Jarvis）', '初始版本，完整商业方案']
for col_idx, val in enumerate(data):
    table.rows[1].cells[col_idx].text = val

# ========== 保存 ==========
output_path = '/root/medi-slim/docs/MediSlim-商业方案-v1.0.docx'
doc.save(output_path)
print(f"✅ Word文档已生成：{output_path}")
print(f"📄 文件大小：{os.path.getsize(output_path) / 1024:.0f} KB")
